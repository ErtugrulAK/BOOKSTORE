import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls

def add_heading_styled(doc, text, level, space_before=12, space_after=6):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(space_before)
    heading.paragraph_format.space_after = Pt(space_after)
    heading.paragraph_format.keep_with_next = True
    
    run = heading.runs[0]
    run.font.name = 'Segoe UI'
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(26, 54, 93)  # Navy Blue
        run.bold = True
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(43, 108, 176)  # Mid Blue
        run.bold = True
    return heading

def add_p_styled(doc, text, space_after=6, bold=False, italic=False, size=10.5, color=RGBColor(45, 55, 72)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return p

def main():
    doc = docx.Document()
    
    # Page Setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Document Header
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(18)
    run_title = p_title.add_run("DOKUZ EYLÜL ÜNİVERSİTESİ MÜHENDİSLİK FAKÜLTESİ\nKİTAP SATIŞ SİSTEMİ YASAL SÖZLEŞMELERİ")
    run_title.font.name = 'Segoe UI'
    run_title.font.size = Pt(18)
    run_title.font.color.rgb = RGBColor(26, 54, 93)
    run_title.bold = True

    p_div = doc.add_paragraph()
    p_div_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="A0AEC0"/></w:pBdr>')
    p_div._p.get_or_add_pPr().append(p_div_border)
    p_div.paragraph_format.space_after = Pt(24)

    # --- SECTION 1: ÖN BİLGİLENDİRME FORMU ---
    add_heading_styled(doc, "BÖLÜM 1: ÖN BİLGİLENDİRME FORMU", level=1)
    
    pre_info_intro = (
        "İşbu Ön Bilgilendirme Formu, Alıcı’nın https://kitapsatis.deu.edu.tr/ internet sitesi üzerinden kitap/yayın satın almadan önce; "
        "ürünün temel nitelikleri, fiyatı, ödeme, teslimat, cayma hakkı ve iade koşulları hakkında bilgilendirilmesi amacıyla hazırlanmıştır."
    )
    add_p_styled(doc, pre_info_intro, space_after=12, italic=True)

    pre_info_sections = [
        ("1. Satıcı Bilgileri", 
         "Ünvan: Dokuz Eylül Üniversitesi Mühendislik Fakültesi Dekanlığı\nAdres: Merkez Yerleşkesi, Buca / İzmir\nE-posta: kitapsatis@deu.edu.tr\nWeb Sitesi: https://kitapsatis.deu.edu.tr/"),
        
        ("2. Sözleşme Konusu Ürün / Hizmet Bilgileri", 
         "Alıcı tarafından satın alınan kitap/yayın ürününün adı, ISBN numarası, adedi, birim fiyatı, KDV dahil toplam tutarı, kargo bedeli ve genel toplamı sipariş tamamlama ekranında ve sipariş onay e-postasında yer almaktadır.\nGüncel fiyatlandırma, KDV oranı ve stok bilgileri Satıcı'nın internet sitesinde ilan edildiği şekilde geçerlidir."),
        
        ("3. Ödeme Bilgileri", 
         "Alıcı, sipariş bedelini kredi kartı veya banka kartı ile sanal POS aracılığıyla ödeyebilir.\nÖn provizyon siparişin verildiği anda alınır; sipariş onayı sonrasında tahsilat gerçekleştirilir."),
        
        ("4. Teslimat Bilgileri", 
         "Teslimat, Alıcı'nın sipariş sırasında bildirdiği adrese anlaşmalı kargo firması aracılığıyla yapılır.\nSatıcı, ödemenin kendisine ulaşmasından itibaren en fazla 30 (otuz) gün içinde teslimatı gerçekleştirmekle yükümlüdür. Yurt içi teslimatlar, kargoya teslim edildiği günden itibaren genellikle 2-5 iş günü içinde tamamlanır; bu süre kargo firmasının hizmet koşullarına göre değişebilir.\nAlıcı, teslimat adresine ilişkin bilgilerin eksiksiz ve doğru olduğunu kabul eder. Hatalı veya eksik adres bilgisi nedeniyle doğabilecek ek kargo masrafları Alıcı'ya aittir."),
        
        ("5. Cayma Hakkı", 
         "Alıcı, sözleşmenin kurulduğu tarihten itibaren 14 (on dört) gün içinde herhangi bir gerekçe göstermeksizin ve cezai şart ödemeksizin cayma hakkını kullanabilir.\nCayma hakkının kullanılabilmesi için ürünün ambalajı açılmamış, kullanılmamış, deforme edilmemiş ve yeniden satılabilir durumda olması gerekir.\nCayma hakkını kullanmak isteyen Alıcı, Satıcı'ya kitapsatis@deu.edu.tr adresi üzerinden e-posta yoluyla veya yazılı olarak bildirimde bulunur.\nAlıcı, cayma bildiriminin ardından ürünü 10 (on) gün içinde, kargo masrafları kendisine ait olmak üzere Satıcı'ya iade eder. Satıcı, iade edilen ürünü teslim aldığı tarihten itibaren 14 (on dört) gün içinde ödemeyi Alıcı'ya iade eder."),
        
        ("6. Cayma Hakkının Kullanılamayacağı Durumlar", 
         "Alıcı tarafından ambalajı açılan, okunmaya başlanan veya üzerinde işaretleme/notlama yapılan kitap ve yayınlarda cayma hakkı kullanılamaz.\nSipariş üzerine özel baskı/cilt yaptırılan veya kişiye özel hazırlanan yayınlarda cayma hakkı kullanılamaz.\nDijital içerik ve e-kitap ürünlerinde, indirme veya aktivasyon işlemi gerçekleştirilmişse cayma hakkı kullanılamamaktadır."),
        
        ("7. Hasarlı, Eksik veya Yanlış Ürün Teslimi", 
         "Teslimat sırasında hasarlı ya da eksik ürün alınması halinde Alıcı, durumu teslim tarihinden itibaren 3 (üç) iş günü içinde fotoğraf ile belgeleyerek Satıcı'ya e-posta yoluyla bildirmelidir.\nSatıcı, inceleme sonucuna göre ürünü değiştirir veya bedelini iade eder. Yanlış ürün gönderilmesi durumunda iade kargo masrafları Satıcı tarafından karşılanır."),
        
        ("8. Kişisel Verilerin Korunması", 
         "Alıcı'nın kişisel verileri, 6698 sayılı Kişisel Verilerin Korunması Kanunu kapsamında işlenmektedir. Satıcı'nın KVKK Aydınlatma Metni internet sitesinde kamuoyuyla paylaşılmıştır."),
        
        ("9. Uyuşmazlıkların Çözümü", 
         "Bu form ve mesafeli satış sözleşmesinden doğabilecek uyuşmazlıklarda İzmir Mahkemeleri, İcra Daireleri ve Tüketici Hakem Heyetleri yetkilidir."),
        
        ("10. Onay", 
         "Alıcı; ürünün temel nitelikleri, satış fiyatı, ödeme şekli, teslimat koşulları, cayma hakkı ve iade şartları hakkında önceden bilgilendirildiğini, işbu Ön Bilgilendirme Formu'nu okup anladığını ve elektronik ortamda onayladığını kabul, beyan ve taahhüt eder.")
    ]

    for title, text in pre_info_sections:
        add_heading_styled(doc, title, level=2)
        add_p_styled(doc, text, space_after=12)

    # Page Break for the next contract
    doc.add_page_break()

    # --- SECTION 2: MESAFELİ SATIŞ SÖZLEŞMESİ ---
    add_heading_styled(doc, "BÖLÜM 2: MESAFELİ SATIŞ SÖZLEŞMESİ", level=1)
    
    mss_sections = [
        ("MADDE 1 – Sözleşmenin Tarafları", 
         "(1) Bu sözleşme; Dokuz Eylül Üniversitesi Mühendislik Fakültesi, Merkez Yerleşkesi, Buca/İzmir adresinde faaliyet gösteren Dokuz Eylül Üniversitesi Mühendislik Fakültesi Kitap Satış Birimi (bundan sonra “Satıcı” olarak anılacaktır.) ile sipariş formunda iletişim bilgileri yer alan kişi (bundan sonra “Alıcı” olarak anılacaktır.) arasında akdedilmiştir.\n(2) Satıcı ve Alıcı ayrı ayrı anıldığında “Taraf”; birlikte anıldığında “Taraflar” ibaresi kullanılacaktır."),
        
        ("MADDE 2 – Sözleşmenin Konusu", 
         "(1) İşbu sözleşmenin konusu; Alıcı'nın, Satıcı'ya ait https://kitapsatis.deu.edu.tr/ internet adresi üzerinden satın aldığı, nitelikleri ve bedeli aşağıda belirtilen kitap/yayın ürününün satışı ve teslimi ile ilgili olarak 6502 sayılı Tüketicinin Korunması Hakkında Kanun ve Mesafeli Sözleşmeler Yönetmeliği hükümleri çerçevesinde Taraflar'ın hak ve yükümlülüklerinin belirlenmesidir."),
        
        ("MADDE 3 – Satıcıya İlişkin Bilgiler", 
         "Ünvan: Dokuz Eylül Üniversitesi Mühendislik Fakültesi\nAdres: Merkez Yerleşkesi, Buca/İzmir\nE-posta: kitapsatis@deu.edu.tr\nWeb sitesi: https://kitapsatis.deu.edu.tr/"),
        
        ("MADDE 4 – Alıcıya İlişkin Bilgiler", 
         "(1) Alıcı'ya ilişkin ad-soyad, teslimat adresi, telefon ve e-posta bilgileri; Alıcı tarafından sipariş aşamasında doldurulan üyelik/sipariş formunda beyan edilen bilgilerdir. Alıcı, söz konusu bilgilerin doğru ve eksiksiz olduğunu kabul ve taahhüt eder; hatalı bilgi nedeniyle doğabilecek her türlü sonuçtan bizzat sorumludur."),
        
        ("MADDE 5 – Sözleşme Konusu Ürün", 
         "(1) Sözleşme konusu ürünün/ürünlerin adı, ISBN numarası, adedi, birim fiyatı (KDV dahil), toplam tutarı ve kargo bedeline ilişkin bilgiler; Alıcı'nın sipariş tamamlama ekranında onayladığı sipariş özetinde ve sipariş onay e-postasında yer almaktadır. Söz konusu sipariş özeti bu sözleşmenin ayrılmaz bir parçasını oluşturur.\n(2) KDV oranı, güncel fiyatlandırma ve stok bilgileri Satıcı'nın internet sitesinde ilan edildiği şekilde geçerlidir."),
        
        ("MADDE 6 – Ödeme ve Sipariş", 
         "(1) Alıcı, sipariş bedelini; kredi/banka kartı ile sanal POS aracılığıyla ödeyebilir.\n(2) Ön provizyon siparişin verildiği anda alınır; sipariş onayı sonrasında tahsilat gerçekleştirilir."),
        
        ("MADDE 7 – Teslimat", 
         "(1) Teslimat, Alıcı'nın sipariş sırasında bildirdiği adrese anlaşmalı kargo firması aracılığıyla yapılır. Satıcı, ödemenin kendisine ulaşmasından itibaren en fazla 30 (otuz) gün içinde teslimatı gerçekleştirmekle yükümlüdür.\n(2) Yurt içi teslimatlar kargoya teslim edildiği günden itibaren genellikle 2–5 iş günü içinde tamamlanır; bu süre kargo firmasının hizmet koşullarına göre değişebilir.\n(3) Alıcı, teslimat adresine ilişkin bilgilerin eksiksiz ve doğru olduğunu taahhüt eder. Hatalı adres bilgisi nedeniyle doğan ek kargo masrafları Alıcı'ya aittir.\n(4) Kargo firmasından kaynaklanan gecikmeler Satıcı'nın sorumluluğunda değildir; Satıcı bu durumlarda Alıcı'yı derhal bilgilendirmekle yükümlüdür."),
        
        ("MADDE 8 – Cayma Hakkı", 
         "(1) Alıcı, sözleşmenin kurulduğu tarihten itibaren 14 (on dört) gün içinde herhangi bir gerekçe göstermeksizin ve cezai şart ödemeksizin cayma hakkını kullanabilir.\n(2) Cayma hakkının kullanılabilmesi için ürünün aynı durum ve koşullarda (ambalajı açılmamış, kullanılmamış, deforme edilmemiş biçimde) iade edilmesi şarttır.\n(3) Aşağıdaki durumlarda cayma hakkı kullanılamaz:\na) Alıcı tarafından ambalajı açılan, okunmaya başlanan veya üzerinde işaretleme/notlama yapılan kitap ve yayınlar.\nb) Sipariş üzerine özel baskı/cilt yaptırılan veya kişiye özel hazırlanan yayınlar.\nc) Dijital içerik ve e-kitap ürünler (indirme/aktivasyon işlemi gerçekleştirilmişse).\n(4) Cayma hakkını kullanmak isteyen Alıcı, Satıcı'ya e-posta veya yazılı bildirimde bulunur. Bildirim tarihinden itibaren 10 (on) gün içinde ürünü, kargo masrafları kendisine ait olmak üzere Satıcı'ya iade eder. Satıcı, iade edilen ürünü teslim aldığı tarihten itibaren 14 (on dört) gün içinde ödemeyi Alıcı'ya iade eder."),
        
        ("MADDE 9 – Hasarlı, Eksik veya Yanlış Ürün", 
         "(1) Teslimat sırasında hasarlı ya da eksik ürün alınması halinde Alıcı, durumu teslim tarihinden itibaren 3 (üç) iş günü içinde fotoğraf ile belgeleyerek Satıcı'ya e-posta yoluyla bildirmelidir. Satıcı, inceleme sonucuna göre ürünü değiştirir veya bedelini iade eder.\n(2) Yanlış ürün gönderilmesi durumunda iade kargo masrafları Satıcı tarafından karşılanır."),
        
        ("MADDE 10 – Genel Hükümler", 
         "(1) Alıcı, sipariş aşamasında ürüne ait temel nitelikler, fiyat ve teslimat koşullarına ilişkin ön bilgileri okuduğunu ve onayladığını kabul, beyan ve taahhüt eder.\n(2) Satıcı, stok tükenmesi veya temin edilemeyen ürün söz konusu olduğunda Alıcı'yı derhal bilgilendirerek tahsil edilen bedeli en geç 14 (on dört) gün içinde iade eder.\n(3) Alıcı, işbu sözleşmeden doğan haklarını ve yükümlülüklerini Satıcı'nın yazılı onayı olmaksızın üçüncü kişilere devredemez ve temlik edemez."),
        
        ("MADDE 11 – Gizlilik ve Kişisel Veriler", 
         "(1) Tarafların her biri, diğer taraftan edindiği kişisel ve ticari bilgileri yalnızca bu sözleşmenin amacı doğrultusunda kullanacak; üçüncü kişilere açıklamayacak ve gizli tutacaktır. Bu yükümlülük sözleşmenin sona ermesinden sonra da geçerliliğini korur.\n(2) Alıcı'nın kişisel verileri, 6698 sayılı Kişisel Verilerin Korunması Kanunu (KVKK) kapsamında işlenmekte olup Satıcı'nın KVKK Aydınlatma Metni internet sitesinde kamuoyuyla paylaşılmıştır."),
        
        ("MADDE 12 – Mücbir Sebepler", 
         "(1) Doğal afet, savaş, terör, salgın hastalık, hükümet kısıtlamaları, yangın, sel veya benzeri tarafların kontrolü dışındaki olaylar nedeniyle yükümlülüklerin yerine getirilememesi halinde taraflar birbirine karşı sorumlu tutulamaz. Mücbir sebebin ortaya çıkması halinde etkilenen taraf, diğer tarafı 3 (üzen) gün içinde yazılı olarak bilgilendirir."),
        
        ("MADDE 13 – Feragat ve Kısmi Geçersizlik", 
         "(1) Herhangi bir tarafın sözleşme hükümlerinden birini uygulamaması, söz konusu haktan feragat olarak yorumlanamaz.\n(2) Bu sözleşmenin herhangi bir hükmünün geçersiz ya da uygulanamaz olması, diğer hükümlerin geçerliliğini etkilemez."),
        
        ("MADDE 14 – Yetkili Yargı Yeri", 
         "(1) Taraflar arasında işbu sözleşmeden doğabilecek ihtilafların çözümünde İzmir Mahkemeleri, İcra Daireleri ve Tüketici Hakem Heyetleri yetkilidir."),
        
        ("MADDE 15 – Sözleşmenin Yürürlüğü", 
         "(1) İşbu sözleşme 15 (on beş) madde ve gerekli eklerinden ibaret olup Alıcı'nın internet sitesi üzerinden siparişini onaylamasından itibaren yürürlüğe girer ve taraflar için bağlayıcı hale gelir.")
    ]

    for title, text in mss_sections:
        add_heading_styled(doc, title, level=2)
        add_p_styled(doc, text, space_after=12)

    # Page Break for the KVKK section
    doc.add_page_break()

    # --- SECTION 3: KVKK AYDINLATMA METNİ ---
    add_heading_styled(doc, "BÖLÜM 3: KİŞİSEL VERİLERİN KORUNMASI KANUNU (KVKK) AYDINLATMA METNİ", level=1)
    
    kvkk_intro = (
        "6698 sayılı Kişisel Verilerin Korunması Kanunu (“KVKK”) uyarınca, kişisel verileriniz; veri sorumlusu olarak "
        "Dokuz Eylül Üniversitesi Mühendislik Fakültesi Dekanlığı (“Fakülte/Dekanlık”) tarafından aşağıda açıklanan kapsamda işlenebilecektir."
    )
    add_p_styled(doc, kvkk_intro, space_after=12)

    kvkk_sections = [
        ("1. Veri Sorumlusu",
         "Veri sorumlusu Dokuz Eylül Üniversitesi Mühendislik Fakültesi Dekanlığı'dır."),
        
        ("2. Kişisel Verilerin Hangi Amaçla İşleneceği",
         "Toplanan kişisel verileriniz (Ad, soyad, T.C. Kimlik / Öğrenci No, teslimat adresi, telefon numarası, e-posta adresi, sipariş ve ödeme yöntemi detayları);\n"
         "• Kitap satış işlemlerinin gerçekleştirilmesi, siparişlerinizin hazırlanması ve faturalandırılması,\n"
         "• Satın alınan kitapların elden teslimi (Dekanlık teslimat kodu doğrulama işlemleri) veya kargo yoluyla ulaştırılması,\n"
         "• Tüketici mevzuatından kaynaklanan yükümlülüklerin yerine getirilmesi,\n"
         "• Yetkili kamu kurum ve kuruluşlarına yasal bilgi sağlama yükümlülüklerinin ifası amaçlarıyla sınırlı olarak işlenmektedir."),
        
        ("3. İşlenen Kişisel Verilerin Kimlere ve Hangi Amaçla Aktarılabileceği",
         "Kişisel verileriniz, yukarıda belirtilen amaçların gerçekleştirilmesi doğrultusunda;\n"
         "• Satın alınan kitapların teslimatı amacıyla anlaşmalı kargo şirketlerine (Örn: PTT Kargo),\n"
         "• Ödemelerin güvenli şekilde tahsil edilmesi amacıyla aracı banka ve sanal POS altyapı sağlayıcılarına (Örn: Ziraat Bankası),\n"
         "• Kanuni yükümlülüklerin ifası kapsamında yetkili kamu kurum ve kuruluşlarına yasal sınırlar dahilinde aktarılabilecektir."),
        
        ("4. Kişisel Veri Toplamanın Yöntemi ve Hukuki Sebebi",
         "Kişisel verileriniz, web sitesine üye olmanız, sipariş formunu doldurmanız ve ödeme adımlarını tamamlamanız esnasında tamamen elektronik ortamda toplanmaktadır. Söz konusu veriler, KVKK Madde 5/2 kapsamında “bir sözleşmenin kurulması veya ifasıyla doğrudan doğruya ilgili olması” ve “veri sorumlusunun hukuki yükümlülüğünü yerine getirebilmesi” hukuki sebeplerine dayanarak işlenmektedir."),
        
        ("5. Veri Sahibinin KVKK Madde 11 Kapsamındaki Hakları",
         "Dilediğiniz zaman Dekanlığımıza başvurarak kişisel verilerinizin; işlenip işlenmediğini öğrenme, işlenme amacına uygun kullanılıp kullanılmadığını bilme, yurt içinde aktarıldığı üçüncü kişileri öğrenme ve eksik/yanlış işlenmişse düzeltilmesini talep etme haklarına sahipsiniz.")
    ]

    for title, text in kvkk_sections:
        add_heading_styled(doc, title, level=2)
        add_p_styled(doc, text, space_after=12)

    # Save to file
    out_filename = "DEÜ Kitap Satış - Yasal Sözleşmeler ve KVKK Metni.docx"
    doc.save(out_filename)
    print(f"Generated successfully: {out_filename}")

if __name__ == "__main__":
    main()
