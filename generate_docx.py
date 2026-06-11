import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def add_heading_styled(doc, text, level, space_before=12, space_after=6):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(space_before)
    heading.paragraph_format.space_after = Pt(space_after)
    heading.paragraph_format.keep_with_next = True
    
    run = heading.runs[0]
    run.font.name = 'Segoe UI'
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(26, 54, 93)  # Navy Blue
        run.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(43, 108, 176)  # Mid Blue
        run.bold = True
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(74, 85, 104)  # Slate Gray
        run.bold = True
    return heading

def main():
    doc = docx.Document()
    
    # Page setup - Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Base Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(45, 55, 72) # Dark Gray for readability
    
    # Title / Header Block
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(6)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("DEÜ KİTAP SATIŞ PLATFORMU")
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(26, 54, 93)
    title_run.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Sanal POS Entegrasyonu Onay ve Uyumluluk Belgesi")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(43, 108, 176)
    sub_run.bold = True
    
    # Info block table (Merchant details for the bank)
    table = doc.add_table(rows=5, cols=2)
    table.alignment = 1 # Center
    table.autofit = False
    
    info_data = [
        ("Üye İşyeri Ticari Unvanı:", "Dokuz Eylül Üniversitesi Mühendislik Fakültesi Dekanlığı"),
        ("Faaliyet Gösterilen Web Adresi:", "http://localhost:5173 (DEÜ Kitap Satış Arayüzü)"),
        ("Entegre Edilen POS Altyapısı:", "Ziraat Bankası Ortak Ödeme Sayfası (VFT Altyapısı)"),
        ("SSL Güvenlik Standardı:", "256-bit SSL / HTTPS Şifreleme Protokolü"),
        ("Belge Hazırlanma Amacı:", "Banka denetim ve POS onay süreci görsel ve teknik uyumluluk doğrulaması")
    ]
    
    # Format table widths
    col_widths = [Inches(2.5), Inches(4.0)]
    for r_idx, (label, val) in enumerate(info_data):
        row = table.rows[r_idx]
        
        # Label cell
        cell_lbl = row.cells[0]
        cell_lbl.width = col_widths[0]
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.paragraph_format.space_before = Pt(4)
        p_lbl.paragraph_format.space_after = Pt(4)
        run_lbl = p_lbl.add_run(label)
        run_lbl.bold = True
        run_lbl.font.size = Pt(9.5)
        run_lbl.font.color.rgb = RGBColor(74, 85, 104)
        
        # Value cell
        cell_val = row.cells[1]
        cell_val.width = col_widths[1]
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.space_before = Pt(4)
        p_val.paragraph_format.space_after = Pt(4)
        run_val = p_val.add_run(val)
        run_val.font.size = Pt(9.5)
    
    # Line divider
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_before = Pt(12)
    p_div_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="A0AEC0"/></w:pBdr>')
    p_div._p.get_or_add_pPr().append(p_div_border)
    
    # 1. Giriş ve Uyumluluk Beyanı
    add_heading_styled(doc, "1. Giriş ve Entegrasyon Beyanı", level=1)
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.line_spacing = 1.15
    p_intro.paragraph_format.space_after = Pt(12)
    p_intro.add_run(
        "Bu doküman, Dokuz Eylül Üniversitesi Kitap Satış Platformu'nun bankacılık düzenlemelerine ve "
        "Virtual POS (Sanal POS) onay kriterlerine olan uyumluluğunu doğrulamak amacıyla hazırlanmıştır. "
        "Sistem üzerinde ödeme akışları güvenli 3D Secure ve Ziraat Bankası ortak ödeme sayfası entegrasyonu standartlarına "
        "göre yapılandırılmış olup, yasal yükümlülükler doğrultusunda gerekli sözleşmeler ve güvenlik beyanları site "
        "arayüzünde konumlandırılmıştır."
    )
    
    # VPOS Check List table
    add_heading_styled(doc, "Banka Sanal POS Denetim Kontrol Listesi", level=2)
    check_table = doc.add_table(rows=6, cols=3)
    check_table.alignment = 1 # Center
    
    # Headers
    hdr_cells = check_table.rows[0].cells
    hdr_cells[0].text = "Denetim Kriteri"
    hdr_cells[1].text = "Sistemdeki Durumu"
    hdr_cells[2].text = "Uyumluluk"
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        
    checklist = [
        ("KDV Dahil Toplam Tutar Gösterimi", "Sepet sayfasında ürün bazlı ve genel toplam KDV dahil net olarak gösterilir.", "✓ UYUMLU"),
        ("Kart Bilgileri Şifreleme Standardı", "Ziraat Bankası güvenli ortak ödeme alt yapısı ve 256-bit SSL güvenlik beyanı.", "✓ UYUMLU"),
        ("Mesafeli Satış Sözleşmesi (MSS)", "Footer grubunda ve tüm sayfalarda erişilebilir şekilde aktif olarak yer alır.", "✓ UYUMLU"),
        ("İptal, İade ve İptal Politikası", "İletişim ve Mesafeli Satış Sözleşmesi sayfalarında iade/iptal prosedürü açıklanmıştır.", "✓ UYUMLU"),
        ("İşyeri İletişim ve Adres Bilgileri", "Tam unvan, açık adres, sabit telefon ve e-posta footer ve iletişimde listelenir.", "✓ UYUMLU"),
    ]
    
    for idx, (crit, desc, status) in enumerate(checklist):
        row = check_table.rows[idx+1]
        row.cells[0].text = crit
        row.cells[1].text = desc
        row.cells[2].text = status
        
        # Format font size
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(9.0)
        row.cells[2].paragraphs[0].runs[0].font.bold = True
        row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(22, 163, 74) # Green for uyumlu
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # 2. Ekran Görüntüleri ve Entegrasyon Aşamaları
    add_heading_styled(doc, "2. Uyumluluk ve Entegrasyon Ekran Görüntüleri", level=1)
    
    compliance_steps = [
        {
            "title": "Alışveriş Sepeti ve Sipariş Özeti",
            "image": "manual_assets/03_cart_step1.png",
            "desc": (
                "Kullanıcının satın almak istediği kitapların adı, kategorisi, adet miktarı ve KDV dahil net birim fiyatı listelenir. "
                "Sağ kısımda yer alan 'Sipariş Özeti' kartında sepet toplam tutarı net bir şekilde gösterilmektedir. "
                "Sanal POS denetimlerinde zorunlu olan 'tutarın ve ürünlerin net gösterimi' şartı bu arayüzde sağlanmaktadır."
            )
        },
        {
            "title": "Teslimat ve Fatura Adresi Formu",
            "image": "manual_assets/04_cart_step2.png",
            "desc": (
                "Siparişin teslim edileceği alıcının adı, soyadı, telefon numarası ve açık adresi bu form aracılığıyla alınır. "
                "İl ve İlçe alanları dinamik ve aramalı seçim listesinden seçilerek hatalı veri girişi engellenir. "
                "Fatura ve teslimat bilgilerinin eksiksiz alınması, ödeme işlemi öncesi doğrulama için zorunludur."
            )
        },
        {
            "title": "Kart Bilgileri Giriş Ekranı ve Güvenlik Beyanı",
            "image": "manual_assets/05_cart_step3.png",
            "desc": (
                "Kullanıcı kart sahibinin adı-soyadı, 16 haneli kart numarası, son kullanma tarihi ve CVV güvenlik kodunu bu güvenli form alanlarına girer. "
                "Kart bilgileri girilirken sayfanın altında banka denetimlerinde kesinlikle bulunması zorunlu olan yeşil çerçeveli "
                "'Güvenli Ödeme Protokolü: Ödemeniz Ziraat Bankası güvenli ortak ödeme sistemi altyapısı üzerinden 256-bit SSL şifreleme protokolü ile korunmaktadır' "
                "güvenlik ibaresi ve SSL kilidi (🔒) simgesi yer almaktadır. Kart verileri kesinlikle yerel sunucularda saklanmamakta, banka VPOS sistemine güvenli kanaldan iletilmektedir."
            )
        },
        {
            "title": "Ödeme Başarı / Sipariş Onay Ekranı",
            "image": "manual_assets/06_cart_step4.png",
            "desc": (
                "Kart sorgulaması banka POS sisteminden olumlu döndüğünde kullanıcıya gösterilen onay ekranıdır. "
                "Ekranda sipariş numarası (#ORD-... şeklinde güncellenmiş), sipariş tarihi ve tahsil edilen toplam tutar "
                "açıkça gösterilmektedir. Bu sayfa, ödeme işleminin başarıyla tamamlandığını ve veritabanı kayıtlarının güncellendiğini teyit eder."
            )
        },
        {
            "title": "Yasal Metinler ve Mesafeli Satış Sözleşmesi",
            "image": "manual_assets/07_distance_sales.png",
            "desc": (
                "Banka POS onay süreçlerinin en kritik şartlarından biri olan 'Mesafeli Satış Sözleşmesi' arayüzüdür. "
                "Sitede yasal mevzuat gereğince hazırlanmış olan sözleşme metni aktif bir sayfa olarak yayınlanmaktadır. "
                "Bu metin, alıcı ve satıcı arasındaki hukuki yükümlülükleri, iade ve cayma hakkı koşullarını detaylandırmaktadır."
            )
        },
        {
            "title": "Kurumsal İletişim ve Üye İşyeri Adres Bilgileri",
            "image": "manual_assets/08_contact.png",
            "desc": (
                "Sanal POS başvurusunun kabul edilmesi için sitenin üye işyerine ait iletişim kanallarını açıkça göstermesi şarttır. "
                "İletişim sayfasında Mühendislik Fakültesi Dekanlığı tam adresi, sabit telefon numarası, faks numarası, "
                "kurumsal e-posta adresi (kitapsatis@deu.edu.tr) ve Google Haritalar konumu yer almaktadır. Aynı iletişim bilgileri "
                "ve kurumsal linkler sitenin alt kısmındaki (footer) sabit menüde de yer alarak sürekli erişilebilir kılınmıştır."
            )
        }
    ]
    
    for idx, step in enumerate(compliance_steps):
        # Heading
        add_heading_styled(doc, f"Ekran {idx+1}: {step['title']}", level=2)
        
        # Description
        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.line_spacing = 1.15
        p_desc.paragraph_format.space_after = Pt(10)
        p_desc.add_run(step['desc'])
        
        # Image
        img_path = step['image']
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_after = Pt(12)
            p_img.add_run().add_picture(img_path, width=Inches(5.2))
            
            # Caption
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(20)
            cap_run = p_cap.add_run(f"Görsel {idx+1}: {step['title']} Arayüz Entegrasyonu")
            cap_run.font.size = Pt(9.0)
            cap_run.font.color.rgb = RGBColor(113, 128, 150)
            cap_run.italic = True
        else:
            p_err = doc.add_paragraph()
            p_err.add_run(f"[Görsel Eksik: {img_path}]").font.color.rgb = RGBColor(220, 38, 38)
            
    # Save document
    filename = "DEÜ Kitap Satış - Ödeme Sistemi Entegrasyon Raporu.docx"
    doc.save(filename)
    print(f"Successfully generated compliance document: {filename}")

if __name__ == "__main__":
    main()
